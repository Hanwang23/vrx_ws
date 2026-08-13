from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/home/han/Ai_ws/Study/vrx_ws")
OUT = ROOT / "group_meeting_ppt"
DOCX_OUT = OUT / "han_usv_controller_本条对话修改工作说明.docx"
ASSETS = OUT / "assets"


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def set_cell_text(cell, text, bold=False, color="24324A", size=10.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Noto Sans CJK SC"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.name = "Noto Sans CJK SC"
    r.font.size = Pt(10.5)
    r.bold = bold
    return p


def add_image(doc, name, width=15.5, caption=None):
    path = ASSETS / name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Cm(width))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(6)
        r = cp.add_run(caption)
        r.font.name = "Noto Sans CJK SC"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(96, 112, 125)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(item)
        r.font.name = "Noto Sans CJK SC"
        r.font.size = Pt(10.2)


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("han_usv_controller 本条对话修改工作说明  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 130, 140)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.7)
    sec.bottom_margin = Cm(1.6)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    add_page_number(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Noto Sans CJK SC"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Noto Sans CJK SC"
    styles["Title"].font.size = Pt(25)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Noto Sans CJK SC"
    styles["Heading 1"].font.color.rgb = RGBColor(23, 50, 77)
    styles["Heading 2"].font.name = "Noto Sans CJK SC"
    styles["Heading 2"].font.color.rgb = RGBColor(36, 127, 120)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    p.add_run().add_picture(str(ASSETS / "sim_title.png"), width=Cm(16.5))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    r = title.add_run("han_usv_controller 本条对话修改工作说明")
    r.font.name = "Noto Sans CJK SC"
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(23, 50, 77)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("VRX WAM-V 自主航行系统：从能运行到可解释、可验收")
    r.font.name = "Noto Sans CJK SC"
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(36, 127, 120)
    add_body(doc, "本文集中说明本条对话中对 han_usv_controller 所做的工作，突出修改范围、实现链路、可视化结果和自动验收证据。")
    doc.add_page_break()

    add_heading(doc, "一、工作目标与总体改造", 1)
    add_body(doc, "本次工作不是只增加一个避障功能，而是围绕 VRX WAM-V 自主航行任务建立一条完整的工程闭环：感知与定位提供稳定状态，滚动地图承载环境记忆，规划与控制负责完成航点，动态目标和 COLREGs 提供会遇监督，RViz 与自动评测负责解释和验收。")
    add_image(doc, "sim_card.png", 15.5, "图 1  仿真场景：工作围绕同一套 WAM-V 场景、控制器和验收脚本展开")
    add_bullets(doc, [
        "统一推荐启动入口：colregs_learning.launch.py。",
        "将算法、场景、RViz、测试和评测脚本纳入同一条运行主链。",
        "把“是否到达”扩展为“为什么这样走、是否安全、是否可重复验证”。",
    ])

    add_heading(doc, "二、具体修改工作量", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["工作模块", "完成的修改内容", "形成的结果"]
    for i, text in enumerate(headers):
        shade(table.rows[0].cells[i], "187F78")
        set_cell_text(table.rows[0].cells[i], text, bold=True, color="FFFFFF", size=10.5)
    rows = [
        ("感知与定位", "接入 robot_localization EKF；增加 PlanarEKF 健康回退；处理三维激光水面点；增加浮标候选聚类与动态目标跟踪。", "定位状态更稳定，感知结果可在 RViz 检查。"),
        ("地图与规划", "建立 100 m 滚动占据栅格；加入障碍衰减、安全膨胀、Dubins State Lattice A*、40 m 子目标和在线重规划。", "静态障碍具备记忆、绕行和 fallback 监控能力。"),
        ("控制与规则", "实现 ILOS + PID 主闭环、曲率预瞄限速、艏摇前馈、安全制动与倒车恢复；加入 CPA/TCPA 和教学型 COLREGs 监督。", "静态避障和动态会遇分层处理。"),
        ("场景与运行", "整合 Wayfinding、WAM-V、12 枚红绿浮标、2 艘动态目标船、中文 RViz 配置和仿真独占锁。", "启动入口统一，重复启动和场景冲突可控。"),
        ("评测与交付", "增加自动评测脚本、硬性门槛、status 输出、源码/配置/模型测试和本次汇报材料。", "结果可量化、可复盘、可回归。"),
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, size=9.5)
    doc.add_paragraph()
    add_image(doc, "lidar_card.png", 9.5, "图 2  激光与障碍感知结果：点云、岸线和障碍信息进入可视化链路")

    add_heading(doc, "三、整体流程", 1)
    add_body(doc, "本次修改形成的运行流程如下：")
    flow = doc.add_table(rows=2, cols=7)
    flow.style = "Table Grid"
    flow_titles = ["需求拆分", "ROS2 集成", "启动场景", "状态估计", "地图规划", "控制监督", "可视化验收"]
    flow_sub = ["定位/地图/规划/控制/规则", "launch + 参数 + 节点", "Wayfinding + WAM-V", "EKF / PlanarEKF", "栅格 + Lattice", "ILOS/PID + COLREGs", "RViz + 自动评测"]
    for i in range(7):
        shade(flow.rows[0].cells[i], "187F78")
        set_cell_text(flow.rows[0].cells[i], flow_titles[i], bold=True, color="FFFFFF", size=8.5)
        set_cell_text(flow.rows[1].cells[i], flow_sub[i], size=8.5)
    add_image(doc, "path_card.png", 12.0, "图 3  路径跟踪结果：规划路径、实际轨迹和横向误差可被同时检查")
    add_body(doc, "关键控制逻辑：静态障碍进入滚动地图和 State Lattice；动态目标经过 alpha-beta 跟踪、CPA/TCPA 计算和会遇分类后，输出右转偏置与降速缩放；最终由安全层和推进器控制完成执行。")

    add_heading(doc, "四、图文可视化与结果证据", 1)
    add_heading(doc, "1. 路径与轨迹", 2)
    add_body(doc, "通过 RViz 同时展示规划路径、实际轨迹、横向误差和目标点，能够解释船舶为何偏转、是否贴线以及是否发生异常绕行。")
    add_heading(doc, "2. 障碍与浮标", 2)
    add_body(doc, "三维激光经过水面滤除和聚类后，形成浮标候选与障碍信息；Marker 采用定点删除，避免旧标记残留和周期性闪烁。")
    add_heading(doc, "3. 动态目标与 COLREGs", 2)
    add_body(doc, "动态船单独显示跟踪框、风险颜色、TCPA、DCPA 和会遇类型。动态目标不直接写成静态地图墙，同时保留原始激光的近距离安全检查。")
    add_image(doc, "lidar_card.png", 7.2, "图 4  感知可视化：让障碍、岸线和候选目标成为可检查证据")
    add_image(doc, "path_card.png", 7.2, "图 5  控制可视化：让规划路径和实际运动结果形成对照")

    doc.add_page_break()
    add_heading(doc, "五、自动验收结果", 1)
    metrics = doc.add_table(rows=1, cols=4)
    metrics.style = "Table Grid"
    for i, text in enumerate(["指标", "结果", "指标", "结果"]):
        shade(metrics.rows[0].cells[i], "2C78A5")
        set_cell_text(metrics.rows[0].cells[i], text, bold=True, color="FFFFFF")
    metric_rows = [("官方任务完成率", "100%", "碰撞次数", "0"), ("官方平均误差", "1.076 m", "最小净空", "5.952 m"), ("动态目标最大跟踪数", "2", "COLREGs 激活样本", "189"), ("State Lattice fallback", "0", "最大规划耗时", "6.151 ms"), ("最大艏摇角速度", "11.244 deg/s", "源码/配置/模型测试", "172 passed")]
    for row in metric_rows:
        cells = metrics.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, bold=(i % 2 == 1), color=("187F78" if i % 2 == 1 else "24324A"), size=10)
    add_body(doc, "自动验收命令：bash han_usv_controller/scripts/run_colregs_learning.sh。硬性检查覆盖三航点完成、0 碰撞、EKF 健康、滚动地图有效、动态目标数量、COLREGs 激活次数和 fallback 状态。", bold=True)

    add_heading(doc, "六、工作结论与边界", 1)
    add_bullets(doc, [
        "体系化：从单脚本控制扩展为包含 launch、配置、模型、算法、脚本和测试的 ROS2 工程包。",
        "可解释：路径、轨迹、误差、占据栅格、浮标候选、动态目标和控制状态均有可视化出口。",
        "可回归：把完成率、碰撞、净空、误差、耗时、动态目标和 COLREGs 激活写成验收门槛。",
        "边界：动态船检测使用仿真专用 Topic；激光不能识别红绿语义；COLREGs 属于教学监督层，不是认证避碰系统。",
    ])
    add_image(doc, "sim_title.png", 13.0, "图 6  本次工作的最终定位：面向 VRX 学习和后续研究的可验证实验平台")

    doc.core_properties.title = "han_usv_controller 本条对话修改工作说明"
    doc.core_properties.author = "HanW"
    doc.save(DOCX_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    build()
